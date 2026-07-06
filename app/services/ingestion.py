"""
Ingestion Service
-----------------
1. Parse uploaded file (PDF or .txt)
2. Chunk text with RecursiveCharacterTextSplitter
3. Embed chunks (sentence-transformers, runs locally — free)
4. Store in ChromaDB (vector search)
5. Build / update BM25 index (keyword search)
"""

import uuid
import time
import pickle
from pathlib import Path
from typing import List, Tuple

import chromadb
from chromadb.config import Settings as ChromaSettings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from rank_bm25 import BM25Okapi
import pypdf

from app.core.config import settings

# ── Singletons (loaded once at startup) ─────────────────────
_chroma_client = None
_collection = None
_embedder: SentenceTransformer | None = None
_bm25_index: BM25Okapi | None = None
_bm25_corpus: List[dict] | None = None   # list of {chunk_id, tokens}
BM25_CACHE = Path("./bm25_index.pkl")


def _get_chroma():
    global _chroma_client, _collection
    if _chroma_client is None:
        _chroma_client = chromadb.PersistentClient(
            path=settings.CHROMA_PERSIST_DIR,
            settings=ChromaSettings(anonymized_telemetry=False),
        )
        _collection = _chroma_client.get_or_create_collection(
            name=settings.CHROMA_COLLECTION,
            metadata={"hnsw:space": "cosine"},
        )
    return _collection


def _get_embedder() -> SentenceTransformer:
    global _embedder
    if _embedder is None:
        _embedder = SentenceTransformer(settings.EMBEDDING_MODEL)
    return _embedder


def _load_bm25():
    global _bm25_index, _bm25_corpus
    if BM25_CACHE.exists():
        with open(BM25_CACHE, "rb") as f:
            data = pickle.load(f)
        _bm25_corpus = data["corpus"]
        _bm25_index  = BM25Okapi([c["tokens"] for c in _bm25_corpus])


def _save_bm25():
    with open(BM25_CACHE, "wb") as f:
        pickle.dump({"corpus": _bm25_corpus}, f)


# ── Text extraction ──────────────────────────────────────────
def _extract_text_pdf(file_bytes: bytes) -> List[Tuple[str, int]]:
    """Returns list of (page_text, page_number)."""
    import io
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append((text, i + 1))
    return pages


def _extract_text_docx(file_bytes: bytes) -> List[Tuple[str, int]]:
    import docx
    import io
    doc = docx.Document(io.BytesIO(file_bytes))
    text_parts = []
    
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text.strip())
            
    for table in doc.tables:
        for row in table.rows:
            row_vals = [cell.text.strip() for cell in row.cells]
            if any(row_vals):
                text_parts.append(" | ".join(row_vals))
                
    full_text = "\n".join(text_parts)
    return [(full_text, 1)]


def _extract_text_xlsx(file_bytes: bytes) -> List[Tuple[str, int]]:
    import openpyxl
    import io
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    text_parts = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_parts.append(f"--- Sheet: {sheet_name} ---")
        for row in sheet.iter_rows(values_only=True):
            if any(row):
                row_str = " | ".join([str(val).strip() if val is not None else "" for val in row])
                text_parts.append(row_str)
                
    full_text = "\n".join(text_parts)
    return [(full_text, 1)]


def _extract_text_csv(file_bytes: bytes) -> List[Tuple[str, int]]:
    import csv
    import io
    content = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(content))
    text_parts = []
    
    for row in reader:
        if any(row):
            text_parts.append(", ".join(row))
            
    full_text = "\n".join(text_parts)
    return [(full_text, 1)]


def _extract_text_plain(file_bytes: bytes) -> List[Tuple[str, int]]:
    return [(file_bytes.decode("utf-8", errors="ignore"), 1)]


# ── Main ingestion function ──────────────────────────────────
async def ingest_document(
    file_bytes: bytes,
    filename: str,
    domain: str = "general",
    doc_id: int = 0,
    # Security metadata — inherited into every chunk
    department: str = "General",
    access_level: str = "internal",
    allowed_roles: list | None = None,
    document_type: str = "general",
    confidentiality_level: str = "low",
) -> int:
    """Ingest a document. Returns number of chunks created."""
    global _bm25_corpus, _bm25_index

    # 1. Extract text
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        pages = _extract_text_pdf(file_bytes)
    elif suffix == ".docx":
        pages = _extract_text_docx(file_bytes)
    elif suffix == ".xlsx":
        pages = _extract_text_xlsx(file_bytes)
    elif suffix == ".csv":
        pages = _extract_text_csv(file_bytes)
    else:
        pages = _extract_text_plain(file_bytes)

    # 2. Chunk
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ".", " "],
    )
    chunks, metadatas, ids = [], [], []
    chunk_index = 0
    for page_text, page_num in pages:
        for chunk_text in splitter.split_text(page_text):
            chunk_id = str(uuid.uuid4())
            chunks.append(chunk_text)
            import json as _json
            metadatas.append({
                "source_file": filename,
                "domain": domain,
                "source_page": page_num,
                "doc_id": doc_id,
                "chunk_index": chunk_index,
                "chunk_id": chunk_id,
                # ── Security metadata (inherited from document) ──
                "department": department,
                "access_level": access_level,
                "allowed_roles": _json.dumps(allowed_roles or []),
                "document_type": document_type,
                "confidentiality_level": confidentiality_level,
            })
            ids.append(chunk_id)
            chunk_index += 1

    if not chunks:
        return 0

    # 3. Embed
    embedder = _get_embedder()
    embeddings = embedder.encode(chunks, show_progress_bar=False).tolist()

    # 4. Store in ChromaDB
    col = _get_chroma()
    col.add(documents=chunks, embeddings=embeddings, metadatas=metadatas, ids=ids)

    # 5. Update BM25 index
    _load_bm25()
    if _bm25_corpus is None:
        _bm25_corpus = []

    for chunk_text, chunk_id in zip(chunks, ids):
        _bm25_corpus.append({
            "chunk_id": chunk_id,
            "tokens": chunk_text.lower().split(),
            "text": chunk_text,
        })

    _bm25_index = BM25Okapi([c["tokens"] for c in _bm25_corpus])
    _save_bm25()

    return len(chunks)


# ── Expose singletons for retrieval service ──────────────────
def get_collection():
    return _get_chroma()

def get_embedder():
    return _get_embedder()

def get_bm25() -> Tuple[BM25Okapi | None, List[dict] | None]:
    _load_bm25()
    return _bm25_index, _bm25_corpus
